from pathlib import Path
from datetime import datetime
import os
import re
#import subprocess

from dotenv import load_dotenv
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from openai import OpenAI
from pydantic import BaseModel
from docx import Document


load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)

app.mount("/generated", StaticFiles(directory=str(OUTPUT_DIR)), name="generated")


class JobData(BaseModel):
    url: str = ""
    title: str = ""
    company: str = ""
    description: str = ""

# def convert_docx_to_pdf(docx_path: Path) -> Path:
#     subprocess.run(
#         [
#             "/Applications/LibreOffice.app/Contents/MacOS/soffice",
#             "--headless",
#             "--convert-to",
#             "pdf",
#             "--outdir",
#             str(docx_path.parent),
#             str(docx_path),
#         ],
#         check=True,
#     )
# 
#     return docx_path.with_suffix(".pdf")

def sanitize_filename_part(text: str) -> str:
    text = text or ""
    text = text.replace("\u00A0", " ")
    text = re.sub(r"[^a-zA-Z0-9_-]+", "_", text)
    text = re.sub(r"_+", "_", text)
    return text.strip("_") or "unknown"


def load_master_resume() -> str:
    resume_file = Path("master_resume.txt")
    if resume_file.exists():
        return resume_file.read_text(encoding="utf-8")
    return "No master_resume.txt found."


def ask_openai(prompt: str) -> str:
    response = client.responses.create(
        model="gpt-4.1-mini",
        input=prompt,
    )
    return response.output_text


def save_docx(path: Path, title: str, content: str):
    doc = Document()
    doc.add_heading(title, level=1)

    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith("- ") or line.startswith("• "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.isupper() and len(line) < 80:
            doc.add_heading(line, level=2)
        else:
            doc.add_paragraph(line)

    doc.save(path)


@app.get("/")
def health_check():
    return {"status": "ok"}

def clean_resume(text: str) -> str:
    if not text:
        return ""

    text = text.replace("Tailored Resume", "").strip()

    unwanted_phrases = [
        "This revised ATS-optimized resume",
        "This ATS-optimized resume",
        "This resume highlights",
        "This tailored resume",
    ]

    for phrase in unwanted_phrases:
        idx = text.find(phrase)
        if idx != -1:
            text = text[:idx]

    return text.strip()

@app.post("/api/job-packet")
def create_job_packet(job: JobData):

    master_resume = load_master_resume()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    safe_company = sanitize_filename_part(job.company)
    safe_title = sanitize_filename_part(job.title)
    print("======== JOB DATA ========")
    print("TITLE:", repr(job.title))
    print("COMPANY:", repr(job.company))
    print("URL:", repr(job.url))
    print("==========================")

    resume_filename = f"{timestamp}_{safe_company}_{safe_title}_resume.docx"
    cover_filename = f"{timestamp}_{safe_company}_{safe_title}_cover_letter.docx"
    analysis_filename = f"{timestamp}_{safe_company}_{safe_title}_match_analysis.docx"
    analysis_path = OUTPUT_DIR / analysis_filename
    resume_path = OUTPUT_DIR / resume_filename
    cover_path = OUTPUT_DIR / cover_filename

    resume_prompt = f"""
You are an expert ATS resume editor.

Your task is NOT to rewrite the entire resume.

Rules:
- Keep employers, dates, titles and achievements.
- Do not invent experience.
- Do not invent metrics.
- Preserve structure.
- Improve ATS keyword alignment.
- Return ONLY the resume.
- Do not include 'Tailored Resume'.
- Do not include commentary.

MASTER RESUME:

{master_resume}

TARGET ROLE:

{job.title}

COMPANY:

{job.company}

JOB DESCRIPTION:

{job.description}
"""

    cover_prompt = f"""
Write a professional cover letter.

Requirements:
- Use the actual company name.
- Use the actual role title.
- Use 'Dear Hiring Team,' if no hiring manager is known.
- Do not use placeholders.
- Maximum 350 words.
- Return only the cover letter.
- Do not invent experience.

MASTER RESUME:

{master_resume}

COMPANY:

{job.company}

ROLE:

{job.title}

JOB DESCRIPTION:

{job.description}
"""

    analysis_prompt = f"""
You are a technical recruiting analyst.

Compare this candidate resume against the job description.

Return ONLY this format:

MATCH SCORE: <0-100>

STRONG MATCHES:
- ...

KEYWORD GAPS:
- ...

RECOMMENDED RESUME CHANGES:
- ...

MASTER RESUME:

{master_resume}

JOB TITLE:

{job.title}

COMPANY:

{job.company}

JOB DESCRIPTION:

{job.description}
"""   
    print("ROLE:", job.title)
    print("COMPANY:", job.company)
    match_analysis = ask_openai(analysis_prompt)
    tailored_resume = clean_resume(
        ask_openai(resume_prompt)
    )

    cover_letter = ask_openai(
        cover_prompt
    )

    print(cover_letter[:500])

    save_docx(
        resume_path,
        "Resume",
        tailored_resume
    )

    save_docx(
        cover_path,
        "Cover Letter",
        cover_letter
    )
    save_docx(
        analysis_path,
        "Match Analysis",
        match_analysis
    )
    
#    analysis_pdf_path = convert_docx_to_pdf(analysis_path)
 #   resume_pdf_path = convert_docx_to_pdf(resume_path)
  #  cover_pdf_path = convert_docx_to_pdf(cover_path)
    
    return {
        "matchAnalysisUrl": f"http://localhost:8000/generated/{analysis_filename}",
   #     "matchAnalysisPdfUrl": f"http://localhost:8000/generated/{analysis_pdf_path.name}",
        "resumeUrl": f"http://localhost:8000/generated/{resume_filename}",
#     "resumePdfUrl": f"http://localhost:8000/generated/{resume_pdf_path.name}",
        "coverLetterUrl": f"http://localhost:8000/generated/{cover_filename}",
     #   "coverLetterPdfUrl": f"http://localhost:8000/generated/{cover_pdf_path.name}",
    }    
     